import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = docx.Document()

# Set page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_title = p_title.add_run("🐎 أخيل — منظومة الخدمات الرسمية\nAKHIL SERVICES — MASTER VERSION")
r_title.bold = True
r_title.font.size = Pt(20)
r_title.font.color.rgb = RGBColor(30, 27, 75)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_sub = p_sub.add_run("المرجع التشغيلي والتسعيري المعتمد لخدمات أخيل الـ 10 • أغسطس 2026")
r_sub.font.size = Pt(11)
r_sub.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph("─" * 50)

# Services list
services = [
    ("1. أخيل اقتصادي — AKHIL ECONOMY 🚗", "قاعدة التسعير: BASE × 1.00", "الخدمة الأساسية للانتقال اليومي، والموجهة للعميل الذي يحتاج إلى رحلة عملية وفق معايير أخيل الأساسية للجودة والتشغيل."),
    ("2. أخيل بلس — AKHIL PLUS 🚘", "قاعدة التسعير: BASE × 1.10", "خدمة بمستوى أعلى من AKHIL ECONOMY من حيث حداثة السيارة ومستوى الراحة وتجربة العميل."),
    ("3. أخيل أعمال — AKHIL BUSINESS 💼", "قاعدة التسعير: BASE × 1.20", "الخدمة المميزة لرجال الأعمال والعملاء الراغبين في أرقى مستوى للسيارة وتجربة رحلة استثنائية."),
    ("4. أخيل برثونة — AKHIL PARTHONA 🌸", "قاعدة التسعير: BASE × 1.00 (بدون أي زيادة)", "خدمة أخيل المخصصة للسيدات والتي تتيح طلب رحلة بقيادة برثونة (سائقة معتمدة). اختيار خدمة برثونة لا يمثل في ذاته سبباً لزيادة سعر الرحلة نهائياً."),
    ("5. أخيل رحلات مجدولة — AKHIL TIME ⏱️", "مظلة الجدولة (AKHIL ONE / ROUTINE / CONTRACT)", "AKHIL ONE: رحلة واحدة مجدولة.\nAKHIL ROUTINE: رحلات متكررة ودورية.\nAKHIL CONTRACT: رحلات تعاقدية مسبقة بسعر ثابت مع تخفيض يصل إلى 20%، ولا تسري عليها الزيادة التفاوضية (+10%)."),
    ("6. أخيل إضافي — AKHIL EXTRA ➕", "زيادة +25% لكل شخص (بحد أدنى +50%)", "مخصصة لنقل عدد ركاب إضافي: شخص إضافي (+50%)، شخصان (+50%)، 3 أشخاص (+75%)، 4 أشخاص (+100%)."),
    ("7. أخيل حمولة — AKHIL CARRY 📦", "قاعدة التسعير: BASE × 2.00", "مخصصة لنقل العميل مع حمولة كبيرة (أثاث، أجهزة، بضائع). الحمولة الأساسية المجانية: شنطة كبيرة + شنطة صغيرة. الحمولة الزائدة المقبولة: +50%. الحمولة الكبيرة: AKHIL CARRY (BASE x 2.00)."),
    ("8. أخيل طرد — AKHIL BOX 📫", "قاعدة التسعير: BASE × 1.00", "نقل وتسليم الطرود والأغراض والمستندات دون راكب بنظام النقل المباشر من نقطة استلام إلى نقطة تسليم محددة."),
    ("9. أخيل سفر — AKHIL TRIP 🛣️", "تسعير السفر بين المحافظات", "خدمة رحلات السفر التي تتجاوز الخطوط الحدودية بين المحافظات. تشمل استراحة سفر مجانية 15 دقيقة لكل 100 كم كاملة، والتوقف الإضافي بطلب العميل بـ 3 جنيه/دقيقة."),
    ("10. أخيل شريك الفعاليات — AKHIL PARTNER OF EVENTS 🎪", "تسعير تعاقدي مخصص للحدث", "الشراكة وتنظيم عمليات النقل للفعاليات والمؤتمرات والمناسبات بأسطول مركبات متكامل وجداول تشغيلية مخصصة.")
]

for title, price, desc in services:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p.add_run(f"\n{title}\n")
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = RGBColor(15, 23, 42)
    
    r2 = p.add_run(f"[{price}]\n")
    r2.bold = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(217, 119, 6)
    
    r3 = p.add_run(desc)
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(51, 65, 85)

doc.save('d:\\3altayer.app\\AKHIL_SERVICES_MASTER_VERSION.docx')
print("Successfully generated AKHIL_SERVICES_MASTER_VERSION.docx")
